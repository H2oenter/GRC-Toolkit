"""
Document Exporter — Complete Version
=====================================
Location: grc_toolkit/utils/document_exporter.py

Export functions for all toolkit outputs:
1. export_gap_assessment_xlsx()   — Gap assessment → Excel workbook
2. export_risk_register_xlsx()    — Risk register → Excel workbook
3. export_evidence_tracker_xlsx() — Evidence tracker → Excel workbook
4. export_gap_assessment_docx()   — Gap assessment → Word document
5. export_policy_docx()           — Policy/procedure → Word document
6. export_risk_register_docx()    — Risk register → Word document

Dependencies:
    pip install python-docx openpyxl
"""

import json
import os
import re
from datetime import datetime

# ──────────────────────────────────────────────────────
# Optional imports — degrade gracefully if not installed
# ──────────────────────────────────────────────────────
try:
    from openpyxl import Workbook
    from openpyxl.styles import (
        Font, PatternFill, Alignment, Border, Side, numbers
    )
    from openpyxl.utils import get_column_letter
    HAS_XLSX = True
except ImportError:
    HAS_XLSX = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


# ──────────────────────────────────────────────────────
# SHARED STYLE CONSTANTS
# ──────────────────────────────────────────────────────

# Excel colors
HEADER_FILL_BLUE = "1F4E79"
HEADER_FILL_RED = "C0392B"
HEADER_FILL_PURPLE = "8E44AD"
HEADER_FILL_GREEN = "1E8449"
HEADER_FILL_ORANGE = "D35400"
HEADER_FONT_WHITE = "FFFFFF"

MATURITY_COLORS = {
    "Fully Implemented": "27AE60",
    "Largely Implemented": "2ECC71",
    "Partially Implemented": "F39C12",
    "Minimally Implemented": "E67E22",
    "Not Implemented": "E74C3C",
    "Not Assessed": "95A5A6",
}

RISK_LEVEL_COLORS = {
    "Critical": "E74C3C",
    "High": "E67E22",
    "Medium": "F39C12",
    "Low": "27AE60",
}

PRIORITY_COLORS = {
    "Critical": "E74C3C",
    "High": "E67E22",
    "Medium": "F39C12",
    "Low": "27AE60",
    "N/A": "95A5A6",
}

EVIDENCE_STATUS_COLORS = {
    "Not Collected": "E74C3C",
    "Requested": "E67E22",
    "In Progress": "F39C12",
    "Collected": "2ECC71",
    "Reviewed": "27AE60",
    "Approved": "1E8449",
    "N/A": "95A5A6",
}

THIN_BORDER = None
if HAS_XLSX:
    _side = Side(style="thin", color="CCCCCC")
    THIN_BORDER = Border(left=_side, right=_side, top=_side, bottom=_side)


# ──────────────────────────────────────────────────────
# SHARED EXCEL HELPERS
# ──────────────────────────────────────────────────────

def _xl_header_row(ws, row: int, headers: list, fill_color: str):
    """Write a styled header row in an Excel worksheet."""
    fill = PatternFill(start_color=fill_color, end_color=fill_color,
                       fill_type="solid")
    font = Font(color=HEADER_FONT_WHITE, bold=True, size=11)
    align = Alignment(horizontal="center", vertical="center",
                      wrap_text=True)

    for col_idx, header_text in enumerate(headers, 1):
        cell = ws.cell(row=row, column=col_idx, value=header_text)
        cell.fill = fill
        cell.font = font
        cell.alignment = align
        if THIN_BORDER:
            cell.border = THIN_BORDER


def _xl_data_cell(ws, row: int, col: int, value,
                  fill_color: str = None, font_color: str = None,
                  bold: bool = False, wrap: bool = True):
    """Write a single styled data cell."""
    cell = ws.cell(row=row, column=col, value=value)
    cell.alignment = Alignment(vertical="top", wrap_text=wrap)
    if THIN_BORDER:
        cell.border = THIN_BORDER
    if fill_color:
        cell.fill = PatternFill(start_color=fill_color,
                                end_color=fill_color,
                                fill_type="solid")
    fnt_kwargs = {}
    if font_color:
        fnt_kwargs["color"] = font_color
    if bold:
        fnt_kwargs["bold"] = True
    if fnt_kwargs:
        cell.font = Font(**fnt_kwargs)
    return cell


def _xl_auto_width(ws, min_width: int = 10, max_width: int = 55):
    """Auto-size columns based on header length (approximate)."""
    for col_cells in ws.columns:
        col_letter = get_column_letter(col_cells[0].column)
        header_len = len(str(col_cells[0].value or ""))
        width = max(min_width, min(header_len + 6, max_width))
        ws.column_dimensions[col_letter].width = width


def _xl_set_widths(ws, widths: dict):
    """Set specific column widths.  widths = {"A": 15, "B": 40, ...}"""
    for letter, w in widths.items():
        ws.column_dimensions[letter].width = w


def _ensure_dir(path: str):
    """Create parent directory if needed."""
    parent = os.path.dirname(path)
    if parent:
        os.makedirs(parent, exist_ok=True)


# ══════════════════════════════════════════════════════
# 1. EXPORT GAP ASSESSMENT → EXCEL
# ══════════════════════════════════════════════════════

def export_gap_assessment_xlsx(results: list, metadata: dict,
                                output_path: str):
    """
    Export gap assessment to a multi-sheet Excel workbook.

    Sheets:
        1. Executive Summary — statistics, maturity distribution
        2. Detailed Results  — every control, color-coded
        3. Priority Actions  — gaps sorted by priority
        4. Scores by Domain  — average scores per function/category
    """
    if not HAS_XLSX:
        print("[ERROR] openpyxl not installed. Run: pip install openpyxl")
        return
    _ensure_dir(output_path)

    wb = Workbook()

    company = metadata.get("company", {})
    fw = metadata.get("framework", "N/A")
    ts = metadata.get("timestamp", datetime.now().isoformat())

    total = len(results)
    assessed = [r for r in results if r.get("score")]
    avg_score = (sum(r["score"] for r in assessed) / len(assessed)
                 if assessed else 0)

    maturity_counts = {}
    priority_counts = {}
    for r in results:
        m = r.get("maturity", "Not Assessed")
        p = r.get("priority", "N/A")
        maturity_counts[m] = maturity_counts.get(m, 0) + 1
        priority_counts[p] = priority_counts.get(p, 0) + 1

    # ── Sheet 1: Executive Summary ──
    ws1 = wb.active
    ws1.title = "Executive Summary"
    ws1.sheet_properties.tabColor = "1F4E79"

    title_font = Font(size=20, bold=True, color="1F4E79")
    sub_font = Font(size=14, bold=True, color="333333")
    label_font = Font(size=11, bold=True)

    ws1["A1"] = "Gap Assessment Report"
    ws1["A1"].font = title_font
    ws1.merge_cells("A1:D1")

    ws1["A3"] = "Company:"
    ws1["A3"].font = label_font
    ws1["B3"] = company.get("name", "N/A")
    ws1["A4"] = "Industry:"
    ws1["A4"].font = label_font
    ws1["B4"] = company.get("industry", "N/A")
    ws1["A5"] = "Framework:"
    ws1["A5"].font = label_font
    ws1["B5"] = fw
    ws1["A6"] = "Assessment Date:"
    ws1["A6"].font = label_font
    ws1["B6"] = ts[:10]
    ws1["A7"] = "Company Size:"
    ws1["A7"].font = label_font
    ws1["B7"] = company.get("size", "N/A")

    row = 9
    ws1.cell(row=row, column=1, value="Assessment Statistics").font = sub_font
    row += 1
    ws1.cell(row=row, column=1, value="Total Controls Assessed:").font = label_font
    ws1.cell(row=row, column=2, value=total)
    row += 1
    ws1.cell(row=row, column=1, value="Average Maturity Score:").font = label_font
    ws1.cell(row=row, column=2, value=round(avg_score, 2))
    ws1.cell(row=row, column=3, value="/ 5.0")
    row += 2

    ws1.cell(row=row, column=1, value="Maturity Distribution").font = sub_font
    row += 1
    _xl_header_row(ws1, row, ["Maturity Level", "Count", "Percentage"],
                   HEADER_FILL_BLUE)
    row += 1
    ordered_mat = ["Fully Implemented", "Largely Implemented",
                   "Partially Implemented", "Minimally Implemented",
                   "Not Implemented", "Not Assessed"]
    for mat in ordered_mat:
        cnt = maturity_counts.get(mat, 0)
        if cnt == 0:
            continue
        color = MATURITY_COLORS.get(mat, "FFFFFF")
        _xl_data_cell(ws1, row, 1, mat, fill_color=color,
                      font_color="FFFFFF" if mat in
                      ["Not Implemented", "Fully Implemented"] else None)
        _xl_data_cell(ws1, row, 2, cnt)
        _xl_data_cell(ws1, row, 3,
                      f"{cnt/total*100:.1f}%" if total else "0%")
        row += 1

    row += 1
    ws1.cell(row=row, column=1, value="Priority Distribution").font = sub_font
    row += 1
    _xl_header_row(ws1, row, ["Priority", "Count"], HEADER_FILL_RED)
    row += 1
    for pri in ["Critical", "High", "Medium", "Low", "N/A"]:
        cnt = priority_counts.get(pri, 0)
        if cnt == 0:
            continue
        color = PRIORITY_COLORS.get(pri, "FFFFFF")
        _xl_data_cell(ws1, row, 1, pri, fill_color=color,
                      font_color="FFFFFF" if pri == "Critical" else None)
        _xl_data_cell(ws1, row, 2, cnt)
        row += 1

    _xl_set_widths(ws1, {"A": 28, "B": 30, "C": 15, "D": 15})

    # ── Sheet 2: Detailed Results ──
    ws2 = wb.create_sheet("Detailed Results")
    ws2.sheet_properties.tabColor = "2E86C1"

    headers = [
        "Control ID", "Function", "Category", "Description",
        "Maturity", "Score", "Current State", "Gap",
        "Recommendations", "Priority", "Estimated Effort"
    ]
    _xl_header_row(ws2, 1, headers, HEADER_FILL_BLUE)

    for ridx, r in enumerate(results, 2):
        _xl_data_cell(ws2, ridx, 1,
                      r.get("control_id", ""))
        _xl_data_cell(ws2, ridx, 2,
                      r.get("function", ""))
        _xl_data_cell(ws2, ridx, 3,
                      r.get("category", ""))
        _xl_data_cell(ws2, ridx, 4,
                      r.get("control_description",
                            r.get("description", "")))

        mat = r.get("maturity", "Not Assessed")
        mat_color = MATURITY_COLORS.get(mat, "FFFFFF")
        needs_white = mat in ["Not Implemented", "Fully Implemented"]
        _xl_data_cell(ws2, ridx, 5, mat, fill_color=mat_color,
                      font_color="FFFFFF" if needs_white else None,
                      bold=True)
        _xl_data_cell(ws2, ridx, 6, r.get("score", ""))

        _xl_data_cell(ws2, ridx, 7,
                      r.get("current_state_assessment",
                            r.get("current_state", "")))
        _xl_data_cell(ws2, ridx, 8,
                      r.get("gap", ""))
        _xl_data_cell(ws2, ridx, 9,
                      r.get("recommendations", ""))

        pri = r.get("priority", "N/A")
        pri_color = PRIORITY_COLORS.get(pri, None)
        _xl_data_cell(ws2, ridx, 10, pri,
                      fill_color=pri_color,
                      font_color="FFFFFF" if pri == "Critical" else None)
        _xl_data_cell(ws2, ridx, 11,
                      r.get("estimated_effort", ""))

    _xl_set_widths(ws2, {
        "A": 14, "B": 16, "C": 22, "D": 42,
        "E": 22, "F": 8, "G": 36, "H": 36,
        "I": 42, "J": 13, "K": 16,
    })
    ws2.auto_filter.ref = f"A1:K{len(results)+1}"
    ws2.freeze_panes = "A2"

    # ── Sheet 3: Priority Actions ──
    ws3 = wb.create_sheet("Priority Actions")
    ws3.sheet_properties.tabColor = "C0392B"

    pri_headers = [
        "Priority", "Control ID", "Category", "Maturity", "Score",
        "Gap", "Recommendation", "Effort"
    ]
    _xl_header_row(ws3, 1, pri_headers, HEADER_FILL_RED)

    priority_order = {"Critical": 0, "High": 1, "Medium": 2, "Low": 3}
    gaps = sorted(
        [r for r in results if r.get("score") and r["score"] < 4],
        key=lambda x: (priority_order.get(x.get("priority", "Low"), 4),
                       x.get("score", 5))
    )
    for ridx, r in enumerate(gaps, 2):
        pri = r.get("priority", "N/A")
        _xl_data_cell(ws3, ridx, 1, pri,
                      fill_color=PRIORITY_COLORS.get(pri),
                      font_color="FFFFFF" if pri == "Critical" else None,
                      bold=True)
        _xl_data_cell(ws3, ridx, 2, r.get("control_id", ""))
        _xl_data_cell(ws3, ridx, 3, r.get("category", ""))
        mat = r.get("maturity", "Not Assessed")
        _xl_data_cell(ws3, ridx, 4, mat,
                      fill_color=MATURITY_COLORS.get(mat))
        _xl_data_cell(ws3, ridx, 5, r.get("score", ""))
        _xl_data_cell(ws3, ridx, 6, r.get("gap", ""))
        _xl_data_cell(ws3, ridx, 7, r.get("recommendations", ""))
        _xl_data_cell(ws3, ridx, 8, r.get("estimated_effort", ""))

    _xl_set_widths(ws3, {
        "A": 13, "B": 14, "C": 22, "D": 22, "E": 8,
        "F": 40, "G": 44, "H": 16,
    })
    ws3.auto_filter.ref = f"A1:H{len(gaps)+1}"
    ws3.freeze_panes = "A2"

    # ── Sheet 4: Scores by Domain ──
    ws4 = wb.create_sheet("Scores by Domain")
    ws4.sheet_properties.tabColor = "1E8449"

    dom_headers = ["Function", "Category", "Controls Assessed",
                   "Average Score", "Min Score", "Max Score",
                   "Fully Implemented", "Not Implemented"]
    _xl_header_row(ws4, 1, dom_headers, HEADER_FILL_GREEN)

    # Group by function → category
    domain_stats = {}
    for r in results:
        key = (r.get("function", "Unknown"),
               r.get("category", "Unknown"))
        if key not in domain_stats:
            domain_stats[key] = {"scores": [], "maturities": []}
        if r.get("score"):
            domain_stats[key]["scores"].append(r["score"])
        domain_stats[key]["maturities"].append(
            r.get("maturity", "Not Assessed"))

    ridx = 2
    for (func, cat), stats in sorted(domain_stats.items()):
        scores = stats["scores"]
        mats = stats["maturities"]
        avg = round(sum(scores)/len(scores), 1) if scores else 0
        _xl_data_cell(ws4, ridx, 1, func)
        _xl_data_cell(ws4, ridx, 2, cat)
        _xl_data_cell(ws4, ridx, 3, len(scores))
        _xl_data_cell(ws4, ridx, 4, avg,
                      fill_color="E74C3C" if avg < 2
                      else "F39C12" if avg < 3.5
                      else "27AE60",
                      font_color="FFFFFF" if avg < 2 else None)
        _xl_data_cell(ws4, ridx, 5, min(scores) if scores else "")
        _xl_data_cell(ws4, ridx, 6, max(scores) if scores else "")
        _xl_data_cell(ws4, ridx, 7,
                      mats.count("Fully Implemented"))
        _xl_data_cell(ws4, ridx, 8,
                      mats.count("Not Implemented"))
        ridx += 1

    _xl_set_widths(ws4, {
        "A": 20, "B": 35, "C": 18, "D": 15,
        "E": 12, "F": 12, "G": 20, "H": 18,
    })
    ws4.freeze_panes = "A2"

    # ── Save ──
    wb.save(output_path)
    print(f"[SAVED] Gap Assessment Excel: {output_path}")
    return output_path


# ══════════════════════════════════════════════════════
# 2. EXPORT RISK REGISTER → EXCEL
# ══════════════════════════════════════════════════════

def export_risk_register_xlsx(risks: list, company_name: str,
                               output_path: str):
    """Export risk register to a formatted Excel workbook."""
    if not HAS_XLSX:
        print("[ERROR] openpyxl not installed.")
        return
    _ensure_dir(output_path)

    wb = Workbook()

    # ── Main Sheet ──
    ws = wb.active
    ws.title = "Risk Register"
    ws.sheet_properties.tabColor = "8E44AD"

    headers = [
        "Risk ID", "Title", "Category", "Description",
        "Threat Source", "Vulnerability",
        "Likelihood", "Impact", "Inherent Score", "Inherent Level",
        "Existing Controls",
        "Residual Likelihood", "Residual Impact",
        "Residual Score", "Residual Level",
        "Treatment", "Treatment Plan",
        "Risk Owner", "Target Date", "Related Controls"
    ]
    _xl_header_row(ws, 1, headers, HEADER_FILL_PURPLE)

    for ridx, risk in enumerate(risks, 2):
        _xl_data_cell(ws, ridx, 1, risk.get("risk_id", ""))
        _xl_data_cell(ws, ridx, 2, risk.get("risk_title", ""), bold=True)
        _xl_data_cell(ws, ridx, 3, risk.get("risk_category", ""))
        _xl_data_cell(ws, ridx, 4, risk.get("risk_description", ""))
        _xl_data_cell(ws, ridx, 5, risk.get("threat_source", ""))
        _xl_data_cell(ws, ridx, 6, risk.get("vulnerability", ""))
        _xl_data_cell(ws, ridx, 7, risk.get("likelihood", ""))
        _xl_data_cell(ws, ridx, 8, risk.get("impact", ""))
        _xl_data_cell(ws, ridx, 9, risk.get("inherent_risk_score", ""))

        lvl = risk.get("inherent_risk_level", "")
        _xl_data_cell(ws, ridx, 10, lvl,
                      fill_color=RISK_LEVEL_COLORS.get(lvl),
                      font_color="FFFFFF" if lvl in ["Critical", "High"]
                      else None, bold=True)

        _xl_data_cell(ws, ridx, 11, risk.get("existing_controls", ""))
        _xl_data_cell(ws, ridx, 12, risk.get("residual_likelihood", ""))
        _xl_data_cell(ws, ridx, 13, risk.get("residual_impact", ""))
        _xl_data_cell(ws, ridx, 14, risk.get("residual_risk_score", ""))

        rlvl = risk.get("residual_risk_level", "")
        _xl_data_cell(ws, ridx, 15, rlvl,
                      fill_color=RISK_LEVEL_COLORS.get(rlvl),
                      font_color="FFFFFF" if rlvl == "Critical" else None)

        _xl_data_cell(ws, ridx, 16, risk.get("risk_treatment", ""))
        _xl_data_cell(ws, ridx, 17, risk.get("treatment_plan", ""))
        _xl_data_cell(ws, ridx, 18, risk.get("risk_owner", ""))
        _xl_data_cell(ws, ridx, 19, risk.get("target_date", ""))

        related = risk.get("related_control_ids", [])
        _xl_data_cell(ws, ridx, 20,
                      ", ".join(related) if isinstance(related, list)
                      else str(related))

    _xl_set_widths(ws, {
        "A": 12, "B": 28, "C": 16, "D": 40,
        "E": 20, "F": 28, "G": 12, "H": 10,
        "I": 14, "J": 14, "K": 32,
        "L": 12, "M": 10, "N": 12, "O": 14,
        "P": 14, "Q": 40, "R": 18, "S": 14, "T": 22,
    })
    ws.auto_filter.ref = f"A1:T{len(risks)+1}"
    ws.freeze_panes = "A2"

    # ── Summary Sheet ──
    ws2 = wb.create_sheet("Risk Summary", 0)
    ws2.sheet_properties.tabColor = "6C3483"

    ws2["A1"] = f"Risk Register — {company_name}"
    ws2["A1"].font = Font(size=18, bold=True, color="8E44AD")
    ws2.merge_cells("A1:D1")
    ws2["A3"] = f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"

    row = 5
    ws2.cell(row=row, column=1, value="Risk Level Summary").font = \
        Font(size=14, bold=True)
    row += 1
    _xl_header_row(ws2, row, ["Risk Level", "Count"], HEADER_FILL_PURPLE)
    row += 1
    level_counts = {}
    for r in risks:
        lvl = r.get("inherent_risk_level", "Unknown")
        level_counts[lvl] = level_counts.get(lvl, 0) + 1
    for lvl in ["Critical", "High", "Medium", "Low"]:
        cnt = level_counts.get(lvl, 0)
        if cnt == 0:
            continue
        _xl_data_cell(ws2, row, 1, lvl,
                      fill_color=RISK_LEVEL_COLORS.get(lvl),
                      font_color="FFFFFF" if lvl in ["Critical"] else None,
                      bold=True)
        _xl_data_cell(ws2, row, 2, cnt)
        row += 1

    row += 1
    ws2.cell(row=row, column=1, value="Treatment Strategy Summary").font = \
        Font(size=14, bold=True)
    row += 1
    _xl_header_row(ws2, row, ["Treatment", "Count"], HEADER_FILL_PURPLE)
    row += 1
    treat_counts = {}
    for r in risks:
        t = r.get("risk_treatment", "Unknown")
        treat_counts[t] = treat_counts.get(t, 0) + 1
    for t, c in sorted(treat_counts.items()):
        _xl_data_cell(ws2, row, 1, t, bold=True)
        _xl_data_cell(ws2, row, 2, c)
        row += 1

    _xl_set_widths(ws2, {"A": 28, "B": 15, "C": 15, "D": 15})

    wb.save(output_path)
    print(f"[SAVED] Risk Register Excel: {output_path}")
    return output_path


# ══════════════════════════════════════════════════════
# 3. EXPORT EVIDENCE TRACKER → EXCEL
# ══════════════════════════════════════════════════════

def export_evidence_tracker_xlsx(evidence_items: list, company_name: str,
                                  framework: str, output_path: str):
    """Export evidence tracker to a formatted Excel workbook."""
    if not HAS_XLSX:
        print("[ERROR] openpyxl not installed.")
        return
    _ensure_dir(output_path)

    wb = Workbook()

    # ── Summary Sheet ──
    ws1 = wb.active
    ws1.title = "Summary"
    ws1.sheet_properties.tabColor = "D35400"

    ws1["A1"] = f"Evidence Tracker — {company_name}"
    ws1["A1"].font = Font(size=18, bold=True, color="D35400")
    ws1.merge_cells("A1:D1")
    ws1["A3"] = f"Framework: {framework}"
    ws1["A3"].font = Font(size=12)
    ws1["A4"] = f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"

    total = len(evidence_items)
    status_counts = {}
    for e in evidence_items:
        s = e.get("status", "Not Collected")
        status_counts[s] = status_counts.get(s, 0) + 1
    collected = sum(v for k, v in status_counts.items()
                    if k in ["Collected", "Reviewed", "Approved"])
    pct = round(collected / total * 100, 1) if total > 0 else 0

    row = 6
    ws1.cell(row=row, column=1,
             value="Collection Progress").font = Font(size=14, bold=True)
    row += 1
    ws1.cell(row=row, column=1, value="Total Items:").font = Font(bold=True)
    ws1.cell(row=row, column=2, value=total)
    row += 1
    ws1.cell(row=row, column=1, value="Collected:").font = Font(bold=True)
    ws1.cell(row=row, column=2, value=collected)
    row += 1
    ws1.cell(row=row, column=1, value="Outstanding:").font = Font(bold=True)
    ws1.cell(row=row, column=2, value=total - collected)
    row += 1
    ws1.cell(row=row, column=1, value="Completion:").font = Font(bold=True)
    ws1.cell(row=row, column=2, value=f"{pct}%")

    row += 2
    ws1.cell(row=row, column=1,
             value="Status Breakdown").font = Font(size=14, bold=True)
    row += 1
    _xl_header_row(ws1, row, ["Status", "Count", "Percentage"],
                   HEADER_FILL_ORANGE)
    row += 1
    for status_name in ["Not Collected", "Requested", "In Progress",
                         "Collected", "Reviewed", "Approved", "N/A"]:
        cnt = status_counts.get(status_name, 0)
        if cnt == 0:
            continue
        color = EVIDENCE_STATUS_COLORS.get(status_name, "FFFFFF")
        _xl_data_cell(ws1, row, 1, status_name, fill_color=color,
                      font_color="FFFFFF" if status_name in
                      ["Not Collected", "Approved"] else None)
        _xl_data_cell(ws1, row, 2, cnt)
        _xl_data_cell(ws1, row, 3,
                      f"{cnt/total*100:.1f}%" if total else "0%")
        row += 1

    _xl_set_widths(ws1, {"A": 25, "B": 15, "C": 15})

    # ── Detail Sheet ──
    ws2 = wb.create_sheet("Evidence Items")
    ws2.sheet_properties.tabColor = "E67E22"

    ev_headers = [
        "Evidence ID", "Control ID", "Evidence Name", "Type",
        "Description", "Typical Source", "Frequency",
        "Priority", "Status", "Collected Date", "File Path", "Notes"
    ]
    _xl_header_row(ws2, 1, ev_headers, HEADER_FILL_ORANGE)

    for ridx, ev in enumerate(evidence_items, 2):
        _xl_data_cell(ws2, ridx, 1, ev.get("evidence_id", ""))
        _xl_data_cell(ws2, ridx, 2, ev.get("control_id", ""))
        _xl_data_cell(ws2, ridx, 3, ev.get("evidence_name", ""), bold=True)
        _xl_data_cell(ws2, ridx, 4, ev.get("evidence_type", ""))
        _xl_data_cell(ws2, ridx, 5, ev.get("description", ""))
        _xl_data_cell(ws2, ridx, 6, ev.get("typical_source", ""))
        _xl_data_cell(ws2, ridx, 7, ev.get("frequency", ""))

        pri = ev.get("priority", "")
        _xl_data_cell(ws2, ridx, 8, pri,
                      fill_color=PRIORITY_COLORS.get(pri))

        status = ev.get("status", "Not Collected")
        _xl_data_cell(ws2, ridx, 9, status,
                      fill_color=EVIDENCE_STATUS_COLORS.get(status),
                      font_color="FFFFFF" if status in
                      ["Not Collected", "Approved"] else None,
                      bold=True)

        _xl_data_cell(ws2, ridx, 10, ev.get("collected_date", ""))
        _xl_data_cell(ws2, ridx, 11, ev.get("file_path", ""))
        _xl_data_cell(ws2, ridx, 12, ev.get("notes", ""))

    _xl_set_widths(ws2, {
        "A": 14, "B": 14, "C": 32, "D": 14,
        "E": 44, "F": 22, "G": 14,
        "H": 13, "I": 16, "J": 18, "K": 25, "L": 30,
    })
    ws2.auto_filter.ref = f"A1:L{len(evidence_items)+1}"
    ws2.freeze_panes = "A2"

    # ── Outstanding Items Sheet ──
    ws3 = wb.create_sheet("Outstanding")
    ws3.sheet_properties.tabColor = "E74C3C"

    outstanding = [e for e in evidence_items
                   if e.get("status") in
                   ["Not Collected", "Requested", "In Progress"]]

    out_headers = ["Evidence ID", "Control ID", "Evidence Name",
                   "Type", "Source", "Priority", "Status"]
    _xl_header_row(ws3, 1, out_headers, HEADER_FILL_RED)

    for ridx, ev in enumerate(outstanding, 2):
        _xl_data_cell(ws3, ridx, 1, ev.get("evidence_id", ""))
        _xl_data_cell(ws3, ridx, 2, ev.get("control_id", ""))
        _xl_data_cell(ws3, ridx, 3, ev.get("evidence_name", ""))
        _xl_data_cell(ws3, ridx, 4, ev.get("evidence_type", ""))
        _xl_data_cell(ws3, ridx, 5, ev.get("typical_source", ""))
        pri = ev.get("priority", "")
        _xl_data_cell(ws3, ridx, 6, pri,
                      fill_color=PRIORITY_COLORS.get(pri))
        status = ev.get("status", "")
        _xl_data_cell(ws3, ridx, 7, status,
                      fill_color=EVIDENCE_STATUS_COLORS.get(status))

    _xl_set_widths(ws3, {
        "A": 14, "B": 14, "C": 35, "D": 14,
        "E": 22, "F": 13, "G": 16,
    })

    wb.save(output_path)
    print(f"[SAVED] Evidence Tracker Excel: {output_path}")
    return output_path


# ══════════════════════════════════════════════════════
# 4. EXPORT GAP ASSESSMENT → WORD
# ══════════════════════════════════════════════════════

def export_gap_assessment_docx(results: list, metadata: dict,
                                executive_summary: str,
                                output_path: str):
    """Export gap assessment to a Word document report."""
    if not HAS_DOCX:
        print("[ERROR] python-docx not installed. Run: pip install python-docx")
        return
    _ensure_dir(output_path)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(10)
    style.font.name = "Calibri"

    company = metadata.get("company", {})
    fw = metadata.get("framework", "N/A")

    # Title Page
    doc.add_paragraph("")
    doc.add_paragraph("")
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("GAP ASSESSMENT REPORT")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"\n{fw}")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f"\nPrepared for: {company.get('name', 'N/A')}\n").font.size = Pt(14)
    info.add_run(f"Industry: {company.get('industry', 'N/A')}\n").font.size = Pt(12)
    info.add_run(f"Date: {metadata.get('timestamp', '')[:10]}\n").font.size = Pt(12)
    info.add_run(f"\nCONFIDENTIAL").font.size = Pt(11)

    doc.add_page_break()

    # Table of Contents placeholder
    doc.add_heading("Table of Contents", level=1)
    doc.add_paragraph("1. Executive Summary")
    doc.add_paragraph("2. Assessment Statistics")
    doc.add_paragraph("3. Detailed Findings")
    doc.add_paragraph("4. Priority Actions")

    doc.add_page_break()

    # Executive Summary
    doc.add_heading("1. Executive Summary", level=1)
    if executive_summary:
        for line in executive_summary.split("\n"):
            line = line.strip()
            if not line:
                continue
            if line.startswith("# "):
                doc.add_heading(line[2:], level=2)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=3)
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style="List Bullet")
            else:
                doc.add_paragraph(line)
    else:
        doc.add_paragraph("Executive summary not generated.")

    doc.add_page_break()

    # Statistics
    doc.add_heading("2. Assessment Statistics", level=1)
    total = len(results)
    assessed = [r for r in results if r.get("score")]
    avg = round(sum(r["score"] for r in assessed) / len(assessed), 1) if assessed else 0

    stats_table = doc.add_table(rows=3, cols=2, style="Light Grid Accent 1")
    stats_data = [
        ("Total Controls Assessed", str(total)),
        ("Average Maturity Score", f"{avg} / 5.0"),
        ("Assessment Date", metadata.get("timestamp", "")[:10]),
    ]
    for i, (label, val) in enumerate(stats_data):
        stats_table.cell(i, 0).text = label
        stats_table.cell(i, 1).text = val
        for cell in stats_table.rows[i].cells:
            for paragraph in cell.paragraphs:
                paragraph.style.font.size = Pt(10)

    doc.add_paragraph("")

    # Maturity table
    doc.add_heading("Maturity Distribution", level=2)
    maturity_counts = {}
    for r in results:
        m = r.get("maturity", "Not Assessed")
        maturity_counts[m] = maturity_counts.get(m, 0) + 1

    mt = doc.add_table(rows=1, cols=3, style="Light Grid Accent 1")
    mt.cell(0, 0).text = "Maturity Level"
    mt.cell(0, 1).text = "Count"
    mt.cell(0, 2).text = "Percentage"
    for m_name in ["Fully Implemented", "Largely Implemented",
                    "Partially Implemented", "Minimally Implemented",
                    "Not Implemented", "Not Assessed"]:
        cnt = maturity_counts.get(m_name, 0)
        if cnt == 0:
            continue
        row = mt.add_row()
        row.cells[0].text = m_name
        row.cells[1].text = str(cnt)
        row.cells[2].text = f"{cnt/total*100:.1f}%"

    doc.add_page_break()

    # Detailed Findings
    doc.add_heading("3. Detailed Findings", level=1)
    current_function = ""
    for r in results:
        func = r.get("function", "Unknown")
        if func != current_function:
            current_function = func
            doc.add_heading(func, level=2)

        ctrl_id = r.get("control_id", "N/A")
        maturity = r.get("maturity", "Not Assessed")
        score = r.get("score", "N/A")

        doc.add_heading(f"{ctrl_id} — {maturity} (Score: {score}/5)",
                        level=3)

        desc = r.get("control_description", r.get("description", ""))
        if desc:
            p = doc.add_paragraph()
            p.add_run("Control: ").bold = True
            p.add_run(desc)

        gap = r.get("gap", "")
        if gap:
            p = doc.add_paragraph()
            p.add_run("Gap: ").bold = True
            p.add_run(gap)

        recs = r.get("recommendations", "")
        if recs:
            p = doc.add_paragraph()
            p.add_run("Recommendations: ").bold = True
            p.add_run(recs)

        pri = r.get("priority", "")
        effort = r.get("estimated_effort", "")
        if pri or effort:
            p = doc.add_paragraph()
            p.add_run(f"Priority: {pri}  |  Effort: {effort}")

    doc.save(output_path)
    print(f"[SAVED] Gap Assessment Word: {output_path}")
    return output_path


# ══════════════════════════════════════════════════════
# 5. EXPORT POLICY → WORD
# ══════════════════════════════════════════════════════

def export_policy_docx(content: str, policy_name: str,
                       company_name: str, output_path: str):
    """
    Convert a Markdown-formatted policy into a styled Word document.
    Handles: # headings, ## sub-headings, bullet points, bold, tables.
    """
    if not HAS_DOCX:
        print("[ERROR] python-docx not installed.")
        return
    _ensure_dir(output_path)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(10)
    style.font.name = "Calibri"

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(policy_name.upper())
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"\n{company_name}\n").font.size = Pt(14)
    sub.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d')}\n").font.size = Pt(10)
    sub.add_run("CONFIDENTIAL").font.size = Pt(10)

    doc.add_page_break()

    # Parse Markdown content
    lines = content.split("\n")
    in_table = False
    table_rows = []

    for line in lines:
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            if in_table and table_rows:
                _write_md_table(doc, table_rows)
                table_rows = []
                in_table = False
            continue

        # Table row detection
        if "|" in stripped and not stripped.startswith("#"):
            # Check if it's a table separator row
            if re.match(r"^\|[\s\-:|]+\|$", stripped):
                in_table = True
                continue
            if stripped.startswith("|"):
                in_table = True
                cells = [c.strip() for c in stripped.split("|")
                         if c.strip()]
                table_rows.append(cells)
                continue

        # If we were in a table and now we're not, flush it
        if in_table and table_rows:
            _write_md_table(doc, table_rows)
            table_rows = []
            in_table = False

        # Headings
        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=4)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)

        # Bullet points
        elif stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:]
            p = doc.add_paragraph(style="List Bullet")
            _add_md_runs(p, text)

        # Numbered list
        elif re.match(r"^\d+[\.\)]\s", stripped):
            text = re.sub(r"^\d+[\.\)]\s", "", stripped)
            p = doc.add_paragraph(style="List Number")
            _add_md_runs(p, text)

        # Horizontal rule
        elif stripped in ["---", "***", "___"]:
            doc.add_paragraph("─" * 60)

        # Regular paragraph
        else:
            p = doc.add_paragraph()
            _add_md_runs(p, stripped)

    # Flush any remaining table
    if table_rows:
        _write_md_table(doc, table_rows)

    doc.save(output_path)
    print(f"[SAVED] Policy Word doc: {output_path}")
    return output_path


def _add_md_runs(paragraph, text: str):
    """Add runs to a paragraph, handling **bold** and *italic* Markdown."""
    # Pattern: **bold** or __bold__
    parts = re.split(r"(\*\*[^*]+\*\*|__[^_]+__)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("__") and part.endswith("__"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # Handle *italic*
            italic_parts = re.split(r"(\*[^*]+\*|_[^_]+_)", part)
            for ip in italic_parts:
                if ((ip.startswith("*") and ip.endswith("*") and
                     not ip.startswith("**")) or
                    (ip.startswith("_") and ip.endswith("_") and
                     not ip.startswith("__"))):
                    run = paragraph.add_run(ip[1:-1])
                    run.italic = True
                else:
                    if ip:
                        paragraph.add_run(ip)


def _write_md_table(doc, rows: list):
    """Write a Markdown-style table to a Word document."""
    if not rows:
        return
    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols,
                          style="Light Grid Accent 1")

    for ridx, row_data in enumerate(rows):
        for cidx, cell_text in enumerate(row_data):
            if cidx < num_cols:
                table.cell(ridx, cidx).text = cell_text

    # Bold the header row
    if rows:
        for cell in table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

    doc.add_paragraph("")  # spacing after table


# ══════════════════════════════════════════════════════
# 6. EXPORT RISK REGISTER → WORD
# ══════════════════════════════════════════════════════

def export_risk_register_docx(risks: list, company_name: str,
                               framework: str, output_path: str):
    """Export risk register as a Word document."""
    if not HAS_DOCX:
        print("[ERROR] python-docx not installed.")
        return
    _ensure_dir(output_path)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(10)
    style.font.name = "Calibri"

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("RISK REGISTER")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x8E, 0x44, 0xAD)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"\n{company_name}\n").font.size = Pt(14)
    sub.add_run(f"Framework: {framework}\n").font.size = Pt(12)
    sub.add_run(f"Date: {datetime.now().strftime('%Y-%m-%d')}\n").font.size = Pt(10)

    doc.add_page_break()

    # Summary
    doc.add_heading("Risk Summary", level=1)
    level_counts = {}
    for r in risks:
        lvl = r.get("inherent_risk_level", "Unknown")
        level_counts[lvl] = level_counts.get(lvl, 0) + 1

    st = doc.add_table(rows=1, cols=2, style="Light Grid Accent 1")
    st.cell(0, 0).text = "Risk Level"
    st.cell(0, 1).text = "Count"
    for lvl in ["Critical", "High", "Medium", "Low"]:
        cnt = level_counts.get(lvl, 0)
        if cnt == 0:
            continue
        row = st.add_row()
        row.cells[0].text = lvl
        row.cells[1].text = str(cnt)

    doc.add_paragraph("")
    doc.add_paragraph(f"Total Risks: {len(risks)}")

    doc.add_page_break()

    # Individual Risks
    doc.add_heading("Risk Details", level=1)
    for risk in risks:
        rid = risk.get("risk_id", "N/A")
        rtitle = risk.get("risk_title", "Untitled")
        lvl = risk.get("inherent_risk_level", "Unknown")

        doc.add_heading(f"{rid} — {rtitle} [{lvl}]", level=2)

        fields = [
            ("Category", risk.get("risk_category", "")),
            ("Description", risk.get("risk_description", "")),
            ("Threat Source", risk.get("threat_source", "")),
            ("Vulnerability", risk.get("vulnerability", "")),
            ("Likelihood", f"{risk.get('likelihood', '')}/5"),
            ("Impact", f"{risk.get('impact', '')}/5"),
            ("Inherent Risk Score", str(risk.get("inherent_risk_score", ""))),
            ("Inherent Risk Level", lvl),
            ("Existing Controls", risk.get("existing_controls", "")),
            ("Residual Likelihood", f"{risk.get('residual_likelihood', '')}/5"),
            ("Residual Impact", f"{risk.get('residual_impact', '')}/5"),
            ("Residual Risk Score", str(risk.get("residual_risk_score", ""))),
            ("Residual Risk Level", risk.get("residual_risk_level", "")),
            ("Treatment Strategy", risk.get("risk_treatment", "")),
            ("Treatment Plan", risk.get("treatment_plan", "")),
            ("Risk Owner", risk.get("risk_owner", "")),
            ("Target Date", risk.get("target_date", "")),
        ]

        rt = doc.add_table(rows=len(fields), cols=2,
                            style="Light Grid Accent 1")
        for i, (label, value) in enumerate(fields):
            rt.cell(i, 0).text = label
            for paragraph in rt.cell(i, 0).paragraphs:
                for run in paragraph.runs:
                    run.bold = True
            rt.cell(i, 1).text = str(value) if value else ""

        related = risk.get("related_control_ids", [])
        if related:
            doc.add_paragraph(
                f"Related Controls: {', '.join(related) if isinstance(related, list) else related}"
            )
        doc.add_paragraph("")  # spacing

    doc.save(output_path)
    print(f"[SAVED] Risk Register Word: {output_path}")
    return output_path